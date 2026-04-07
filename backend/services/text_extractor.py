"""
Text extraction service for PDF and DOCX files.
Handles text preprocessing and cleaning.
"""
import re
import logging
from pathlib import Path
from typing import Tuple, Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> Tuple[str, dict]:
    """
    Extract text from a PDF file using pdfplumber.
    Returns (text, metadata).
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")

    text_parts = []
    metadata = {"pages": 0, "has_tables": False, "has_images": False}

    with pdfplumber.open(file_path) as pdf:
        metadata["pages"] = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            if page.extract_tables():
                metadata["has_tables"] = True

    full_text = "\n".join(text_parts)
    return clean_text(full_text), metadata


def extract_text_from_docx(file_path: str) -> Tuple[str, dict]:
    """
    Extract text from a DOCX file using python-docx.
    Returns (text, metadata).
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document(file_path)
    metadata = {
        "paragraphs": len(doc.paragraphs),
        "has_tables": len(doc.tables) > 0,
        "sections": len(doc.sections)
    }

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    full_text = "\n".join(text_parts)
    return clean_text(full_text), metadata


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    - Remove excessive whitespace
    - Fix common OCR/extraction artifacts
    - Normalize quotes and dashes
    """
    if not text:
        return ""

    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Remove page numbers (common patterns)
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)

    # Remove excessive blank lines (more than 2 consecutive)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Fix hyphenated line breaks (common in PDFs)
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)

    # Normalize smart quotes to straight quotes
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2018', "'").replace('\u2019', "'")

    # Normalize dashes
    text = text.replace('\u2013', '-').replace('\u2014', '--')

    # Remove null bytes and control characters (except newlines/tabs)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # Collapse multiple spaces to single space
    text = re.sub(r' {2,}', ' ', text)

    return text.strip()


def split_into_sentences(text: str) -> list[str]:
    """
    Split text into sentences using a robust approach.
    Handles academic text with abbreviations, citations, etc.
    """
    # Academic abbreviations that shouldn't trigger sentence splits
    abbreviations = r'(?:Dr|Mr|Mrs|Ms|Prof|Sr|Jr|vs|etc|al|Fig|fig|No|Vol|vol|pp|cf|e\.g|i\.e|et al|Ph\.D|M\.D|B\.Sc|M\.Sc|approx|dept|est|inc|corp|ltd)'

    # Temporarily protect abbreviations
    protected = re.sub(
        rf'\b({abbreviations})\.',
        lambda m: m.group().replace('.', '<<DOT>>'),
        text
    )

    # Protect decimal numbers like 3.14
    protected = re.sub(r'(\d+)\.(\d+)', r'\1<<DOT>>\2', protected)

    # Protect citations like [1], [2,3], (Smith et al., 2023)
    protected = re.sub(r'\[[\d,\s]+\]', '<<CITE>>', protected)

    # Split on sentence boundaries: ". ", "? ", "! " followed by uppercase
    sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z"\(])', protected)

    # Restore protected dots
    sentences = [s.replace('<<DOT>>', '.').replace('<<CITE>>', '') for s in sentences]

    # Filter out very short fragments (likely artifacts)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10]

    return sentences


def detect_sections(text: str) -> list[dict]:
    """
    Detect common academic paper sections.
    Returns list of {name, start_char, end_char}.
    """
    section_patterns = [
        r'\b(abstract)\b',
        r'\b(introduction)\b',
        r'\b(background)\b',
        r'\b(literature review)\b',
        r'\b(methodology|methods|materials and methods)\b',
        r'\b(results|findings)\b',
        r'\b(discussion)\b',
        r'\b(conclusion|conclusions)\b',
        r'\b(references|bibliography)\b',
        r'\b(acknowledgements?|acknowledgments?)\b',
        r'\b(appendix|appendices)\b'
    ]

    sections = []
    text_lower = text.lower()

    for pattern in section_patterns:
        matches = list(re.finditer(pattern, text_lower))
        for match in matches:
            # Verify it looks like a section header (near start of line)
            start = match.start()
            line_start = text_lower.rfind('\n', 0, start) + 1
            pre_text = text_lower[line_start:start].strip()
            if len(pre_text) < 5:  # Near start of line
                sections.append({
                    "name": match.group().title(),
                    "start_char": start,
                    "end_char": match.end()
                })

    # Sort by position
    sections.sort(key=lambda x: x["start_char"])

    # Remove duplicates (keep first occurrence)
    seen_names = set()
    unique_sections = []
    for s in sections:
        name = s["name"].lower()
        if name not in seen_names:
            seen_names.add(name)
            unique_sections.append(s)

    return unique_sections


def count_words(text: str) -> int:
    """Count words in text."""
    return len(text.split())


def extract_text(file_path: str, file_extension: str) -> Tuple[str, dict]:
    """
    Main entry point for text extraction.
    Dispatches to appropriate extractor based on file type.
    """
    ext = file_extension.lower().lstrip('.')

    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ('docx', 'doc'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: pdf, docx")
